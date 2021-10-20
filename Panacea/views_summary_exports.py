#####
# This file contains the views needed to generate pdf reports from html, using WeasyPrint
#####
import base64
import io
import os
import tempfile
from zipfile import ZipFile

from django.contrib.auth.decorators import login_required
from django.core.exceptions import ObjectDoesNotExist
from django.http import HttpResponse, HttpResponseRedirect
from django.template.loader import render_to_string
from docx.image.exceptions import UnrecognizedImageError
from openpyxl.cell import Cell
from weasyprint import HTML
from django.shortcuts import render, redirect

from Panacea.decorators import group_required
from Panacea.utilities import find_user_organization, find_summary_organizations, get_current_summary_report_year
from .builders import ExportReport
from .forms import report_generating_form, create_new_summary_report_form, \
    summary_report_subpart_form, export_organization_select_data, export_organization_select_coversheet
from .models import cover_sheet, tax_rates, organization, report_summary_table, report_summary_table_subpart
from .report_builders import run_reports
from .summary_report_tables_v3 import ReportSummaryTable
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side, Alignment, numbers


@login_required(login_url='/Panacea/login')
@group_required('Summary reporter', 'WSDOT staff')
def cover_sheet_report(request, file_output=None):
    user_org = find_user_organization(request.user.id)
    org_cover_sheet = cover_sheet.objects.get(organization_id=user_org.id)
    agency_type = user_org.summary_organization_classifications.name
    # print(agency_type)
    try:
        base64_logo = base64.encodebytes(org_cover_sheet.organization_logo).decode("utf-8")
    except:
        base64_logo = ""
    if agency_type != "Community provider" or agency_type != "Medicaid broker":
        tax_description, created = tax_rates.objects.get_or_create(organization_id=user_org.id)
        tax_description = tax_description.tax_rate_description
    if file_output == 'pdf':
        html_report = render_to_string('../templates/reports/cover_sheet.html', {'cover_sheet': org_cover_sheet,
                                                                                 'agency_type': agency_type,
                                                                                 'organization': user_org,
                                                                                 'tax_description': tax_description,
                                                                                 'base64_logo': base64_logo,
                                                                                 'file_type': file_output})
        response = HttpResponse(content_type='application/pdf;')
        response['Content-Disposition'] = 'inline; filename=coversheet.pdf'
        response['Content-Transfer-Encoding'] = 'binary'
        HTML(string=html_report).write_pdf(response)

        return response
    else:
        return render(request, '../templates/reports/cover_sheet.html', {'cover_sheet': org_cover_sheet,
                                                                         'agency_type': agency_type,
                                                                         'organization': user_org,
                                                                         'tax_description': tax_description,
                                                                         'base64_logo': base64_logo,
                                                                         'file_type': file_output
                                                                         })

def generate_cover_sheet_report(org_id, file_type='pdf'):
    org = organization.objects.get(id=org_id)
    org_cover_sheet = cover_sheet.objects.get(organization_id=org.id)
    agency_type = org.summary_organization_classifications.name
    # print(agency_type)
    try:
        base64_logo = base64.encodebytes(org_cover_sheet.organization_logo).decode("utf-8")
    except:
        base64_logo = ""
    if agency_type != "Community provider" or agency_type != "Medicaid broker":
        tax_description, created = tax_rates.objects.get_or_create(organization_id=org.id)
        tax_description = tax_description.tax_rate_description

    html_report = render_to_string('../templates/reports/cover_sheet.html', {'cover_sheet': org_cover_sheet,
                                                                             'agency_type': agency_type,
                                                                             'organization': org,
                                                                             'tax_description': tax_description,
                                                                             'base64_logo': base64_logo,
                                                                             'file_type': file_type})
    if file_type == 'pdf':
        pdf_bytes = HTML(string=html_report).write_pdf()
        file = io.BytesIO()
        file.write(pdf_bytes)
        file.seek(0)
    if file_type == 'docx':
        from htmldocx import HtmlToDocx
        from docx import Document
        doc = Document()
        if org_cover_sheet.organization_logo:
            image = io.BytesIO(org_cover_sheet.organization_logo)
            try:
                doc.add_picture(image)
            except UnrecognizedImageError:
                doc.add_paragraph('[Invalid Image Type]')
            except ZeroDivisionError:
                doc.add_paragraph('[Invalid Image Dimensions]')

            image.close()
        new_parser = HtmlToDocx()
        doc_content = new_parser.parse_html_string(html_report)

        for element in doc_content.element.body:
            doc.element.body.append(element)
        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
    file.name = org.name + "." + file_type

    return file


def save_cover_sheet_file(file, path='./cover_sheets/'):
    f = open(os.path.join(path + file.name), 'wb')
    f.write(file.read())
    f.close()


def test_cover_sheet(org_id, file_type='pdf'):
    file = generate_cover_sheet_report(org_id, file_type)
    save_cover_sheet_file(file)


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def generate_all_coversheets_backend_process_not_for_ui(start_with_org_name=None):
    if start_with_org_name:
        skip = True
    else:
        skip = False

    for org in organization.objects.all():
        # print(org.name)
        if org.name == start_with_org_name:
            skip = False
        if not skip:
            if org.name not in ['Washington State Department of Transportation',
                                'Lower Elwha Klallam Tribe',
                                'Nooksack Indian Tribe',
                                'Quileute Nation',
                                'Colville Confederated Tribes',
                                'Heckman Motors, Inc']:
                file, file_type, org = generate_cover_sheet_report(org.id, file_type='docx')
                save_cover_sheet_file(file, file_type, org)

@login_required(login_url='/Panacea/login')
def run_statewide_report_tables(request):
    if request.method == 'POST':
        form = report_generating_form(request.POST)
        if form.is_valid():
            report_list = request.POST.getlist('report_selection')
            # print(report_list)
            size = request.POST.get('report_size')
            run_reports(report_list, size)
    else:
        form = report_generating_form
    return render(request, 'pages/summary/admin/exports/run_statewide_report_tables.html', {'form': form})


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def exports_home(request):
    return render(request, 'pages/summary/admin/exports/exports_home.html', {})

@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def exports_cover_sheets_for_report(request):
    if request.POST:

        org_list = request.POST.getlist('organizations')
        cover_sheet_format = request.POST.get('cover_sheet_format')
        if cover_sheet_format == 'Word':
            file_type = 'docx'
        elif cover_sheet_format == "PDF":
            file_type = 'pdf'
        else:
            raise NotImplementedError

        if type(org_list) != list:
            org_list = [org_list]

        zip_file = io.BytesIO()
        zip_archive = ZipFile(zip_file, 'w')

        for org in org_list:
            # print(org)
            try:
                file = generate_cover_sheet_report(org, file_type=file_type)
                with zip_archive.open(file.name, 'w') as this_file:
                    this_file.write(file.getvalue())
                this_file.close()
            except ObjectDoesNotExist:
                missing_org = organization.objects.get(id=org)
                with zip_archive.open("DOES_NOT_EXIST - " + missing_org.name + "." + file_type, 'w') as this_file:
                    this_file.close()

            file.close()
        zip_archive.close()

        response = HttpResponse(zip_file.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = 'attachment; filename="%s"' % 'coversheets.zip'
        return response


    else:
        form = export_organization_select_coversheet()

        return render(request, 'pages/summary/admin/exports/export_cover_sheets.html', {'form': form})


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def export_data_tables(request, summary_organization_classifications_id=None):
    if request.POST:
        org_list = request.POST.getlist('organizations')
        include_comments = request.POST.get('include_comments')
        # print(include_comments)
        if type(org_list) != list:
            org_list = [org_list]
        # print(org_list)
        zip_file = io.BytesIO()
        zip_archive = ZipFile(zip_file, 'w')

        for org in org_list:
            this_org = organization.objects.get(id=org)
            report_builder = ExportReport(this_org)
            excel_file = report_builder.generate_excel_summary_report(include_comment=include_comments)

            file = io.BytesIO()
            excel_file.save(file)
            with zip_archive.open(this_org.name + ".xlsx", 'w') as this_file:
                this_file.write(file.getvalue())
            this_file.close()
            file.close()
        zip_archive.close()

        response = HttpResponse(zip_file.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = 'attachment; filename="%s"' % 'data_tables.zip'
        return response

    else:
        form = export_organization_select_data()
        if summary_organization_classifications_id:
            from django import forms
            form.fields['organizations'].queryset = find_summary_organizations(summary_organization_classifications_id=summary_organization_classifications_id)

    return render(request, 'pages/summary/admin/exports/export_data_tables.html', {'form': form})


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def review_summary_tables(request):
    report_summary_table_data = report_summary_table.objects.all()
    return render(request, 'pages/summary/admin/exports/review_summary_tables.html',
                  {'report_summary_table_data': report_summary_table_data})

@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def create_new_summary_tables(request):
    form = create_new_summary_report_form()
    if request.POST:
        form = create_new_summary_report_form(request.POST)
        if form.is_valid():
            form.save()
            return redirect(review_summary_tables)
    else:
        form = create_new_summary_report_form()
    return render(request, 'pages/summary/admin/exports/create_new_summary_table.html',
                  {'form': form})


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def edit_summary_tables(request, summary_table_id):
    this_summary_table = report_summary_table.objects.get(id=summary_table_id)
    this_summary_table_sub_parts = this_summary_table.table_sub_part_list.order_by('order')
    form = create_new_summary_report_form(request.POST or None, instance=this_summary_table)
    table_sub_parts = report_summary_table_subpart.objects.all()

    if request.POST and form.is_valid():
        form.save()

    return render(request, 'pages/summary/admin/exports/edit_summary_table.html', {'form': form,
                                                                                   'this_summary_table': this_summary_table,
                                                                                   'this_summary_table_sub_parts': this_summary_table_sub_parts,
                                                                                   'table_sub_parts': table_sub_parts})

@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def edit_create_subpart(request, sub_part_id=None):
    if sub_part_id:
        this_sub_part = report_summary_table_subpart.objects.get(id=sub_part_id)
        form = summary_report_subpart_form(request.POST or None, instance=this_sub_part)
    else:
        form = summary_report_subpart_form(request.POST or None)
    if request.POST:
        # print('post')
        if form.is_valid():
            # print('valid')
            form.save()

    return render(request, 'pages/summary/admin/exports/sub_part_page.html', {'form': form,
                                                                              'sub_part_id': sub_part_id
                                                                              })


def styled_cells(ws, data, first_row_not_number=True):
    i = 0
    for c in data:
        if i == 0 and first_row_not_number:
            c = Cell(ws, column="A", row=1, value=c)
        else:
            if isinstance(c, str):
                if c[0] == '$':
                    c = Cell(ws, column="A", row=1, value=int(c[1:]))
                    c.number_format = numbers.FORMAT_CURRENCY_USD
                elif c[-1] == '%':
                    c = Cell(ws, column="A", row=1, value=float(c[:-1])/100)
                    c.number_format = numbers.FORMAT_PERCENTAGE_00
                elif c.isnumeric():
                    c = float(c)
                    if c.is_integer():
                        c = int(c)
                else:
                    c = Cell(ws, column="A", row=1, value=c)
            if not isinstance(c, Cell):
                if isinstance(c, int):
                    my_format = numbers.BUILTIN_FORMATS[3]
                elif isinstance(c, float):
                    if c.is_integer():
                        my_format = numbers.BUILTIN_FORMATS[3]
                    else:
                        my_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
                else:
                    my_format = numbers.FORMAT_GENERAl
                c = Cell(ws, column="A", row=1, value=c)
                c.number_format = my_format
        i = i + 1
        yield c
    return data


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def get_summary_table(request, summary_table_id):
    this_summary_table = report_summary_table.objects.get(id=summary_table_id)
    report = ReportSummaryTable(this_summary_table)
    report_table = report.produce_table()

    wb = Workbook()
    ws = wb.active

    heading_row = []

    if report.report_summary_table_type == 'Standard':
        heading_row = [report.table_heading]
        report_year = get_current_summary_report_year()
        for i in range(report_year - report.number_of_years_to_pull + 1, report_year + 1):
            heading_row.append(i)
        if report.table_has_percentage_change:
            heading_row.append('One year change (%)')
        ws.append(heading_row)
    elif report.report_summary_table_type == 'Non-standard headings':
        i = 0
        for subpart in report_table:
            if i == 0:
                heading_row = subpart[0]
            subpart = subpart[1:]
            report_table[i] = subpart
            i = i + 1
        ws.append(heading_row)
    else:
        raise NotImplementedError
    # print(report_table)
    for subpart in report_table:
        for row in subpart:
                ws.append(styled_cells(ws, row))
    file = io.BytesIO()
    wb.save(file)
    response = HttpResponse(file.getvalue(), content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename="%s"' % report.save_name + ".xlsx"

    return response


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def add_sub_part_to_table(request, summary_table_id, sub_part_id):
    this_summary_table = report_summary_table.objects.get(id=summary_table_id)
    this_sub_part = report_summary_table_subpart.objects.get(id=sub_part_id)
    this_summary_table.table_sub_part_list.add(this_sub_part)
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def remove_sub_part_from_table(request, summary_table_id, sub_part_id):
    this_summary_table = report_summary_table.objects.get(id=summary_table_id)
    this_sub_part = report_summary_table_subpart.objects.get(id=sub_part_id)
    this_summary_table.table_sub_part_list.remove(this_sub_part)
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


@login_required(login_url='/Panacea/login')
@group_required('WSDOT staff')
def delete_summary_table(request, summary_table_id):
    this_summary_table = report_summary_table.objects.get(id=summary_table_id)
    this_summary_table.delete()
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
